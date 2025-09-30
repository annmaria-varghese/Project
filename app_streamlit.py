import streamlit as st
import fitz  # PyMuPDF
from docx import Document
import os
import tempfile
from io import BytesIO

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain
from langchain_community.llms import HuggingFacePipeline
from langchain.prompts import PromptTemplate

from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline

# ===== Streamlit page config =====
st.set_page_config(page_title="PaperTalk", page_icon="📄", layout="wide")

# ===== Custom CSS for chat bubbles =====
st.markdown("""
<style>
.chat-bubble-user {
    background-color: #DCF8C6;
    border-radius: 12px;
    padding: 10px;
    margin: 5px 0;
    text-align: right;
}
.chat-bubble-bot {
    background-color: #F1F0F0;
    border-radius: 12px;
    padding: 10px;
    margin: 5px 0;
    text-align: left;
}
.stTextInput>div>div>input {
    border-radius: 10px;
    border: 2px solid #4CAF50;
}
</style>
""", unsafe_allow_html=True)

st.title("📄 PaperTalk – Chat with Your Files")
st.markdown("Upload a **PDF/DOCX/TXT/MD** file and start asking questions.")

# ===== File extraction =====
def extract_text(uploaded_file):
    """Read text from PDF, DOCX, TXT, MD"""
    filename = uploaded_file.name
    _, ext = os.path.splitext(filename)
    ext = ext.lower()

    if ext == ".pdf":
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        text = "\n".join(page.get_text("text") for page in doc)

    elif ext == ".docx":
        file_bytes = BytesIO(uploaded_file.read())
        doc = Document(file_bytes)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        text = "\n".join(parts)

    elif ext in [".txt", ".md"]:
        text = uploaded_file.read().decode("utf-8", errors="ignore")

    else:
        raise ValueError(f"Unsupported file type: {ext}")

    # Clean empty lines
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())

# ===== Retriever =====
def get_file_retriever(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
    docs = splitter.create_documents([text])
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vectorstore = FAISS.from_documents(docs, embeddings)
    return vectorstore.as_retriever(search_kwargs={"k": 5})

# ===== Load LLM =====
@st.cache_resource
def load_llm():
    model_name = "google/flan-t5-base"   # ✅ lightweight, free HuggingFace model
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)

    pipe = pipeline(
        task="text2text-generation",
        model=model,
        tokenizer=tokenizer,
        max_length=500
    )
    return HuggingFacePipeline(pipeline=pipe)

# ===== Prompt =====
QA_PROMPT = PromptTemplate(
    input_variables=["context", "question"],
    template="""You are PaperTalk, an AI assistant. 
Answer questions strictly using the provided document context. 
If the answer is not in the document, say: "I couldn’t find that in the document."

Context:
{context}

Question:
{question}

Answer:"""
)

# ===== Main App =====
uploaded_file = st.file_uploader("Upload your file", type=["pdf", "docx", "txt", "md"])

if uploaded_file:
    # Extract + prepare retriever
    text = extract_text(uploaded_file)
    retriever = get_file_retriever(text)
    llm = load_llm()

    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        combine_docs_chain_kwargs={"prompt": QA_PROMPT},
        return_source_documents=True
    )

    if "history" not in st.session_state:
        st.session_state.history = []

    query = st.text_input("💬 Ask me anything about your file:")
    if query:
        response = conversation_chain.invoke({
            "question": query,
            "chat_history": st.session_state.history
        })
        st.session_state.history.append(
            (query, response["answer"], response["source_documents"])
        )

    # Display conversation
    for q, a, srcs in st.session_state.history:
        st.markdown(f"<div class='chat-bubble-user'><b>You:</b> {q}</div>", unsafe_allow_html=True)
        st.markdown(f"<div class='chat-bubble-bot'>🤖 <b>PaperTalk:</b> {a}</div>", unsafe_allow_html=True)

        with st.expander("📖 Sources used for this answer"):
            for i, doc in enumerate(srcs):
                st.markdown(f"**Source {i+1}:** {doc.page_content[:400]}...")
