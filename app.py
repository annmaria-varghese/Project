import streamlit as st
import wikipediaapi
import spacy
import random
import requests
from gtts import gTTS
from io import BytesIO
from docx import Document
from fpdf import FPDF
import networkx as nx
from pyvis.network import Network
import tempfile
import base64
import streamlit.components.v1 as components

# ------------------- App Setup -------------------
st.set_page_config(page_title="QuickThink", layout="wide")

# ------------------- Olive-Cream CSS -------------------
st.markdown(
    """
    <style>
    .stApp { background-color: #f7f6f1; color: #2b2b2b; font-family: "Segoe UI", sans-serif; }
    .title { text-align: center; font-size: 40px; font-weight: 700; color: #3c5a3c !important; margin-bottom: 18px; }
    .card { background-color: #fff; padding: 18px; border-radius: 12px; border-left: 6px solid #4a5d23; margin-bottom: 18px; box-shadow: 0 6px 18px rgba(0,0,0,0.06);}
    .card h3 { color: #4a5d23 !important; margin-bottom: 8px; }
    .fact { padding: 10px 14px; margin: 8px 0; background: #fbfaf6; border-left: 4px solid #4a5d23; border-radius: 6px; font-size: 15px; color: #222; }
    .stButton>button { background-color: #4a5d23; color: #fff; border-radius: 6px; padding: 8px 16px; border: none; }
    .stButton>button:hover { background-color: #3c4d1d; transform: translateY(-2px); }
    .stTextInput>div>div>input { background-color: #fff; color: #222; border-radius: 6px; padding: 10px; border: 1px solid #ddd; }
    </style>
    """,
    unsafe_allow_html=True
)

st.markdown('<div class="title">QuickThink</div>', unsafe_allow_html=True)

# ------------------- Wikipedia & spaCy -------------------
nlp = spacy.load("en_core_web_sm")
wiki = wikipediaapi.Wikipedia(language='en', user_agent='QuickThinkApp/1.0')

# ------------------- Utilities -------------------
def get_online_summary(keyword, max_sentences=10):
    """Fetch summary from Wikipedia REST summary endpoint and return up to max_sentences."""
    url = f"https://en.wikipedia.org/api/rest_v1/page/summary/{keyword}"
    try:
        res = requests.get(url, timeout=6)
        if res.status_code == 200:
            text = res.json().get("extract", "")
            # simple sentence split (keeps punctuation)
            doc = nlp(text)
            sents = [s.text.strip() for s in doc.sents]
            if not sents:
                return text
            return " ".join(sents[:max_sentences]) + (" ..." if len(sents) > max_sentences else "")
    except Exception:
        return None

def get_offline_summary(keyword, max_sentences=10):
    page = wiki.page(keyword)
    if page.exists():
        text = page.text
        doc = nlp(text)
        sents = [s.text.strip() for s in doc.sents if len(s.text.strip()) > 20]
        return " ".join(sents[:max_sentences]) + (" ..." if len(sents) > max_sentences else "")
    return None

def get_random_summary(max_sentences=10):
    """Get a random wiki summary via REST random title endpoint"""
    url = "https://en.wikipedia.org/api/rest_v1/page/random/summary"
    try:
        res = requests.get(url, timeout=6)
        if res.status_code == 200:
            data = res.json()
            title = data.get("title", "Random Topic")
            extract = data.get("extract", "")
            doc = nlp(extract)
            sents = [s.text.strip() for s in doc.sents]
            summary = " ".join(sents[:max_sentences]) + (" ..." if len(sents) > max_sentences else "")
            return title, summary
    except:
        return None, None

def extract_key_takeaways(text, n=3):
    """Select n sentences with most named entities (simple heuristic for 'important')"""
    doc = nlp(text)
    sentences = [sent for sent in doc.sents if len(sent.text.strip()) > 20]
    # score sentence by number of entities it contains + length
    scored = []
    for sent in sentences:
        ents = len([ent for ent in sent.ents])
        scored.append((ents, len(sent.text), sent.text.strip()))
    # sort by (entities desc, length desc)
    scored.sort(key=lambda x: (x[0], x[1]), reverse=True)
    takeaways = [s[2] for s in scored[:n]]
    # fallback: first n sentences
    if len(takeaways) < n:
        takeaways = [sent.text.strip() for sent in sentences[:n]]
    return takeaways

def extract_fun_facts(text, num_facts=4):
    doc = nlp(text)
    sentences = [sent.text.strip() for sent in doc.sents if len(sent.text.strip()) > 30]
    if not sentences:
        return ["No additional insights available."]
    return random.sample(sentences, min(num_facts, len(sentences)))

def generate_quiz_cloze(text, num_q=3):
    """Generate simple cloze MCQ questions by blanking named entities"""
    doc = nlp(text)
    sentences = [sent for sent in doc.sents if len(sent.text.strip()) > 40]
    ents = [ent.text for ent in doc.ents]
    quiz = []
    tried = 0
    i = 0
    while len(quiz) < num_q and tried < 50 and i < len(sentences):
        sent = sentences[i]
        i += 1
        sent_ents = [ent for ent in sent.ents]
        if not sent_ents:
            tried += 1
            continue
        ent = random.choice(sent_ents)
        correct = ent.text
        # build options: correct + up to 3 other entity distractors (unique)
        distractors = [e for e in ents if e.lower() != correct.lower()]
        random.shuffle(distractors)
        options = [correct]
        for d in distractors[:3]:
            options.append(d)
        # ensure we have 4 options; if not, fill with plausible nonsense words from other sentences
        if len(options) < 4:
            filler = [w.text for w in doc if w.is_alpha and len(w.text) > 3]
            random.shuffle(filler)
            for f in filler:
                if f not in options:
                    options.append(f)
                if len(options) == 4:
                    break
        random.shuffle(options)
        question_text = sent.text.replace(correct, "_____")
        quiz.append({"question": question_text, "answer": correct, "options": options})
        tried += 1
    return quiz

def build_entity_graph(text, notebook=False):
    """Return HTML of pyvis graph created from co-occurrence of named entities."""
    doc = nlp(text)
    # collect entities per sentence
    sent_ents = []
    for sent in doc.sents:
        ents = [ent.text for ent in sent.ents if len(ent.text) > 1]
        if len(ents) > 0:
            sent_ents.append(list(dict.fromkeys(ents)))  # unique in sentence preserving order

    G = nx.Graph()
    # add nodes
    for ents in sent_ents:
        for e in ents:
            G.add_node(e)
        # add edges for co-occurrence (all pairs)
        for i in range(len(ents)):
            for j in range(i+1, len(ents)):
                a, b = ents[i], ents[j]
                if G.has_edge(a, b):
                    G[a][b]['weight'] += 1
                else:
                    G.add_edge(a, b, weight=1)

    # build pyvis network
    net = Network(height="420px", width="100%", notebook=notebook)
    net.from_nx(G)
    # small physics
    net.force_atlas_2based()
    html_path = tempfile.NamedTemporaryFile(delete=False, suffix=".html").name
    net.save_graph(html_path)
    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()
    return html

def tts_audio_bytes(text, lang="en"):
    """Return mp3 bytes of TTS rendered text using gTTS."""
    tts = gTTS(text=text, lang=lang, slow=False)
    mp3_fp = BytesIO()
    tts.write_to_fp(mp3_fp)
    mp3_fp.seek(0)
    return mp3_fp.read()

def make_docx(title, summary, takeaways, facts):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(summary)
    doc.add_heading("Key Takeaways", level=2)
    for t in takeaways:
        doc.add_paragraph(t, style='List Bullet')
    doc.add_heading("Interesting Facts", level=2)
    for f in facts:
        doc.add_paragraph(f, style='List Number')
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def make_pdf(title, summary, takeaways, facts):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.set_text_color(43,43,43)
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, title, ln=True)
    pdf.set_font("Arial", size=11)
    pdf.multi_cell(0, 7, summary)
    pdf.ln(4)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 7, "Key Takeaways", ln=True)
    pdf.set_font("Arial", size=11)
    for t in takeaways:
        pdf.multi_cell(0, 7, f"- {t}")
    pdf.ln(2)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 7, "Interesting Facts", ln=True)
    pdf.set_font("Arial", size=11)
    for f in facts:
        pdf.multi_cell(0, 7, f"- {f}")
    out = BytesIO()
    pdf.output(out)
    out.seek(0)
    return out.getvalue()

def download_button_bytes(data: bytes, filename: str, label: str, mime: str):
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:{mime};base64,{b64}" download="{filename}">{label}</a>'
    return href

# ------------------- Sidebar / Controls -------------------
with st.sidebar:
    st.markdown('<div class="card"><h3>QuickThink Features</h3>'
                '<p>Use the controls below to experiment with features.</p></div>', unsafe_allow_html=True)
    # session bookmarks
    if 'bookmarks' not in st.session_state:
        st.session_state['bookmarks'] = {}

# ------------------- Main UI -------------------
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown('<div class="card"><h3>Search Topic</h3></div>', unsafe_allow_html=True)
    keyword = st.text_input("Enter a keyword or phrase", placeholder="e.g., Artificial Intelligence")
    st.write("")  # spacing

    row1c1, row1c2 = st.columns([2,1])
    with row1c1:
        if st.button("Generate"):
            mode = "generate"
            do_random = False
        else:
            mode = None
    with row1c2:
        if st.button("Surprise Me"):
            mode = "surprise"
            do_random = True

    # feature toggles
    st.markdown('<div class="card"><h3>Tools</h3></div>', unsafe_allow_html=True)
    tts_enabled = st.checkbox("Enable Listen Mode (Text-to-Speech)", value=True)
    quiz_enabled = st.checkbox("Enable Quiz (auto generate)", value=True)
    graph_enabled = st.checkbox("Show Knowledge Map", value=True)
    export_enabled = st.checkbox("Show Export Options", value=True)

with col2:
    st.markdown('<div class="card"><h3>Saved Topics</h3></div>', unsafe_allow_html=True)
    if st.session_state['bookmarks']:
        for t, v in st.session_state['bookmarks'].items():
            st.markdown(f"- **{t}**")
    else:
        st.info("No saved topics yet. Use the 'Save Topic' button after generating results.")

# ------------------- Processing -------------------
if 'mode' not in st.session_state:
    st.session_state.mode = None

if 'last_title' not in st.session_state:
    st.session_state.last_title = ""

# Trigger handling
if (st.button("Generate") or st.button("Surprise Me")) or (st.session_state.mode is None and keyword):
    # The above double buttons avoid missing clicks; prefer direct handling below.
    pass

# Execute based on which button was pressed (we inspect Streamlit's formless buttons state)
# Simpler: call both handlers based on explicit checks:
if st.session_state.get("last_action") is None:
    st.session_state["last_action"] = ""

# Determine action by presence of "Surprise Me" click or Generate click:
# Because Streamlit buttons re-run, check the latest click by reading query params isn't trivial.
# We'll do simpler logic: if user clicked Surprise Me (do_random True) or they clicked Generate (mode variable above).

# If user clicked Surprise Me:
if 'Surprise Me' in st.session_state.get('_rerun', ''):
    pass  # ignore; fallback to below behavior

# We'll rely on variables created earlier:
# If do_random set to True via button click above
try:
    do_random
except NameError:
    do_random = False

if do_random:
    title, essay = get_random_summary(max_sentences=10)
    if not essay:
        st.error("Couldn't fetch a random topic. Try again.")
else:
    # only generate when Generate clicked and keyword present
    if (st.button("Generate") or keyword.strip() != ""):
        # We should run when user has entered a keyword and pressed Generate — but due to Streamlit statelessness,
        # we'll attempt to fetch when keyword provided and user pressed the Generate button above earlier.
        # To avoid double network calls, only fetch if keyword non-empty and last_title differs.
        if keyword and (keyword != st.session_state.get('last_title')):
            essay = get_online_summary(keyword, max_sentences=10)
            if not essay:
                essay = get_offline_summary(keyword, max_sentences=10)
            title = keyword
            st.session_state['last_title'] = keyword
        else:
            # either no keyword or same as last one - attempt to use last stored result if present
            essay = None
            title = keyword

# Final guard: if essay not defined yet and keyword exists, try one last time:
if 'essay' not in locals():
    essay = None
if not essay and keyword:
    essay = get_online_summary(keyword, max_sentences=10)
    if not essay:
        essay = get_offline_summary(keyword, max_sentences=10)
    title = keyword

# If still no essay, don't proceed
if not essay:
    st.warning("Enter a keyword and press Generate (or click Surprise Me).")
else:
    # Display results
    st.markdown('<div class="card"><h3>Summary</h3></div>', unsafe_allow_html=True)
    st.write(essay)

    # Key Takeaways
    takeaways = extract_key_takeaways(essay, n=3)
    st.markdown('<div class="card"><h3>Key Takeaways</h3></div>', unsafe_allow_html=True)
    for t in takeaways:
        st.markdown(f'<div class="fact">• {t}</div>', unsafe_allow_html=True)

    # Interesting Facts
    facts = extract_fun_facts(essay, num_facts=4)
    st.markdown('<div class="card"><h3>Interesting Facts</h3></div>', unsafe_allow_html=True)
    for i, f in enumerate(facts, 1):
        st.markdown(f'<div class="fact">{i}. {f}</div>', unsafe_allow_html=True)

    # Text-to-Speech
    if tts_enabled:
        st.markdown('<div class="card"><h3>Listen</h3></div>', unsafe_allow_html=True)
        try:
            audio_bytes = tts_audio_bytes(essay, lang="en")
            st.audio(audio_bytes, format='audio/mp3')
            # provide download link
            href = download_button_bytes(audio_bytes, f"{title}_summary.mp3", "Download audio", "audio/mpeg")
            st.markdown(href, unsafe_allow_html=True)
        except Exception as e:
            st.error("Text-to-speech failed (gTTS might not be available in your environment).")

    # Quiz generation
    if quiz_enabled:
        st.markdown('<div class="card"><h3>Quick Quiz</h3></div>', unsafe_allow_html=True)
        quiz = generate_quiz_cloze(essay, num_q=3)
        if not quiz:
            st.info("Not enough content to make a quiz for this topic.")
        else:
            for q in quiz:
                st.markdown(f"**Q:** {q['question']}")
                options = q['options']
                choice = st.radio("Choose an answer:", options, key=q['question'] + str(random.random()))
                if st.button("Check Answer", key="check_" + q['answer'] + str(random.random())):
                    if choice == q['answer']:
                        st.success("Correct!")
                    else:
                        st.error(f"Incorrect — correct answer: {q['answer']}")

    # Knowledge map
    if graph_enabled:
        st.markdown('<div class="card"><h3>Knowledge Map</h3></div>', unsafe_allow_html=True)
        try:
            html = build_entity_graph(essay, notebook=False)
            components.html(html, height=450, scrolling=True)
        except Exception as e:
            st.info("Could not build graph. Possibly not enough named entities in this topic.")

    # Export options
    if export_enabled:
        st.markdown('<div class="card"><h3>Export & Save</h3></div>', unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1:
            if st.button("Save Topic"):
                st.session_state['bookmarks'][title] = {"summary": essay}
                st.success(f"Saved '{title}' to bookmarks.")
        with c2:
            # DOCX
            docx_bytes = make_docx(title, essay, takeaways, facts)
            b64_docx = base64.b64encode(docx_bytes).decode()
            href_docx = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx}" download="{title}.docx">Download DOCX</a>'
            st.markdown(href_docx, unsafe_allow_html=True)
        with c3:
            # PDF
            pdf_bytes = make_pdf(title, essay, takeaways, facts)
            b64_pdf = base64.b64encode(pdf_bytes).decode()
            href_pdf = f'<a href="data:application/pdf;base64,{b64_pdf}" download="{title}.pdf">Download PDF</a>'
            st.markdown(href_pdf, unsafe_allow_html=True)

    st.markdown('<div class="card"><h3>Explore More</h3><p>Try "Surprise Me" for random topics, or save interesting topics for later.</p></div>', unsafe_allow_html=True)
